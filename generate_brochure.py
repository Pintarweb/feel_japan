import os
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

output_dir = r"d:\AntiGravity Project\feeljapan_brochure\public\2026 Draft Pricing"
os.makedirs(output_dir, exist_ok=True)

# Generate DOCX
doc = Document()
doc.add_heading('MOUNT FUJI CLIMBING 5D4N', 0)
doc.add_paragraph('(Beginning of June - Beginning of Sept)', style='Subtitle')

doc.add_heading('DAY 1: NARITA AIRPORT (L/D)', level=1)
p = doc.add_paragraph()
p.add_run('• Meet and greet by English-speaking tour guide at Narita Airport Lobby.\n')
p.add_run('• Transfer to hotel.\n')
p.add_run('• Stop at Premium Outlet near Narita Airport / Akihabara.\n')
p.add_run('• Lunch at Akihabara.\n')
p.add_run('• Preparation for climbing (purchase climbing accessories).\n')
p.add_run('• Dinner at a restaurant nearby the hotel.\n')
note = p.add_run('(Note: Tour guide will advise what to buy/prepare for breakfast and lunch for the next day. Lunch and dinner at a local restaurant.)')
note.italic = True

doc.add_heading('DAY 2: TRANSFER TO MOUNT FUJI (B, D-Bento)', level=1)
p = doc.add_paragraph()
p.add_run('• Pick up and transfer to Tokyo station.\n')
p.add_run('• Take Shinkansen from Tokyo to Mishima (09:03 → 09:45).\n')
p.add_run('• Meet with the Climbing Team at Mishima North Exit at 10:10 AM and take a bus to Fujinomiya 5th Station.\n')
p.add_run('• 11:40 AM: Arrive at 5th Station, organize equipment, and distribute rental supplies.\n')
p.add_run('• 12:10 PM: Hike to 6th Station mountain hut (approx. 1 hour for altitude acclimatization) & lunch.\n')
p.add_run('• 13:10 PM: Begin the climb.\n')
p.add_run('• 17:00 PM: Arrive at Yamaguchi Sanso (mountain hut).\n')
p.add_run('• 17:30 PM: Dinner and prepare for the next day\'s hike.\n')
p.add_run('• 20:00 PM: Lights out and rest.\n')
note = p.add_run('(Note: The hut is basic, without a bathroom, with shared sleeping areas. Bring earplugs and an eye mask.)')
note.italic = True

doc.add_heading('DAY 3: SUMMIT & DESCENT', level=1)
p = doc.add_paragraph()
p.add_run('• 01:00 AM: Wake up.\n')
p.add_run('• 01:30 AM: Depart for the summit attempt.\n')
p.add_run('• 04:00 AM: Arrive at the summit and enjoy a simple breakfast.\n')
p.add_run('• 04:30-05:00 AM: Sunrise viewing.\n')
p.add_run('• 08:30 AM: Begin the descent.\n')
p.add_run('• 12:00-14:00 PM: Return to Fujinomiya 5th Station, lunch.\n')
p.add_run('• 15:00 PM: Travel to Mishima natural hot spring "Fugetsu no Yu" for a relaxing soak.\n')
p.add_run('• 18:00 PM: Dinner.\n')
note = p.add_run('(Notes: Summit temp is ~0℃. Bring a down jacket, headlamp, gloves, and hat. Protect knees during descent.)')
note.italic = True

doc.add_heading('DAY 4: BACK TO TOKYO (B / L / D)', level=1)
p = doc.add_paragraph()
p.add_run('• Check out of Mishima Onsen Hotel.\n')
p.add_run('• Tokyo tour with guide.\n')
p.add_run('• Breakfast, lunch, and dinner at local restaurants.\n')

doc.add_heading('DAY 5: DEPARTURE (B)', level=1)
p = doc.add_paragraph()
p.add_run('• Check out of the hotel.\n')
p.add_run('• Transfer to Airport / End of service.\n')

doc.add_page_break()

doc.add_heading('Quotation', 0)
doc.add_heading('YIDA Co. Ltd', level=1)
doc.add_paragraph('Osaka, Japan', style='Subtitle')

doc.add_heading('2026 Beginning of June - Beginning of Sept Quotation', level=2)
p = doc.add_paragraph()
p.add_run('• 4 Pax: 326,000 JPY per person (7 seats Alfa)\n').bold = True
p.add_run('• 6 Pax: 274,000 JPY per person (10 seats Hiace)\n').bold = True
p.add_run('• 8 Pax: 248,000 JPY per person (10 seats Hiace)\n').bold = True
p.add_run('(Single supplement: 15,000 JPY per night, per person)').italic = True

doc.add_heading('Price Includes:', level=2)
p = doc.add_paragraph()
p.add_run('• 3 hotel nights with breakfast (Tokyo: Hotel JAL City Tokyo Toyosu, Mishima Onsen: Fugaku no Yu).\n')
p.add_run('• 1 night at the 6th station mountain hut.\n')
p.add_run('• 4 days bus usage + 2 transfers limit to 10 hours and 300km per day.\n')
p.add_run('• Meals: 4x Buffet breakfast, 3x Halal lunch (2,500 JPY), 3x Halal dinner (3,500 JPY), 1x Bento dinner.\n')
p.add_run('(A 10% fee applies if dietary changes made after booking).\n').italic = True
p.add_run('• English-speaking guide for 5 days (Malaysian speaker +8,000 JPY/day).\n')
p.add_run('• Driver and guide\'s accommodation.\n')

doc.add_heading('Price Excludes:', level=2)
p = doc.add_paragraph()
p.add_run('• Bottled water (150 JPY per bottle on the bus).\n')
p.add_run('• Weekend & Holiday surcharges (2,000-6,000 JPY/pax/day).\n')

doc.save(os.path.join(output_dir, 'Fuji_Climbing_Brochure_2026.docx'))

# Generate PDF using fpdf2
class PDF(FPDF):
    def header(self):
        self.set_font('helvetica', 'B', 15)
        self.cell(0, 10, 'MOUNT FUJI CLIMBING 5D4N', align='C', border=0, new_x='LMARGIN', new_y='NEXT')
        self.set_font('helvetica', 'I', 11)
        self.cell(0, 10, '(Beginning of June - Beginning of Sept)', align='C', border=0, new_x='LMARGIN', new_y='NEXT')
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('helvetica', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', align='C')

pdf = PDF()
pdf.add_page()
pdf.set_font("helvetica", size=11)

sections = [
    ("DAY 1: NARITA AIRPORT (L/D)", [
        "- Meet and greet by English-speaking tour guide at Narita Airport Lobby.",
        "- Transfer to hotel.",
        "- Stop at Premium Outlet near Narita Airport / Akihabara.",
        "- Lunch at Akihabara.",
        "- Preparation for climbing (purchase climbing accessories).",
        "- Dinner at a restaurant nearby the hotel.",
        "(Note: Tour guide will advise what to buy/prepare for breakfast and lunch for the next day. Lunch and dinner at a local restaurant.)"
    ]),
    ("DAY 2: TRANSFER TO MOUNT FUJI (B, D-Bento)", [
        "- Pick up and transfer to Tokyo station.",
        "- Take Shinkansen from Tokyo to Mishima (09:03 -> 09:45).",
        "- Meet with the Climbing Team at Mishima North Exit at 10:10 AM and take a bus to Fujinomiya 5th Station.",
        "- 11:40 AM: Arrive at 5th Station, organize equipment, and distribute rental supplies.",
        "- 12:10 PM: Hike to 6th Station mountain hut (approx. 1 hour for altitude acclimatization) & lunch.",
        "- 13:10 PM: Begin the climb.",
        "- 17:00 PM: Arrive at Yamaguchi Sanso (mountain hut).",
        "- 17:30 PM: Dinner and prepare for the next day's hike.",
        "- 20:00 PM: Lights out and rest.",
        "(Note: The hut is basic, without a bathroom, with shared sleeping areas. Bring earplugs and an eye mask.)"
    ]),
    ("DAY 3: SUMMIT & DESCENT", [
        "- 01:00 AM: Wake up.",
        "- 01:30 AM: Depart for the summit attempt.",
        "- 04:00 AM: Arrive at the summit and enjoy a simple breakfast.",
        "- 04:30-05:00 AM: Sunrise viewing.",
        "- 08:30 AM: Begin the descent.",
        "- 12:00-14:00 PM: Return to Fujinomiya 5th Station, lunch.",
        "- 15:00 PM: Travel to Mishima natural hot spring \"Fugetsu no Yu\" for a relaxing soak.",
        "- 18:00 PM: Dinner.",
        "(Notes: Summit temp is ~0 degree C. Bring a down jacket, headlamp, gloves, and hat. Protect knees during descent.)"
    ]),
    ("DAY 4: BACK TO TOKYO (B / L / D)", [
        "- Check out of Mishima Onsen Hotel.",
        "- Tokyo tour with guide.",
        "- Breakfast, lunch, and dinner at local restaurants."
    ]),
    ("DAY 5: DEPARTURE (B)", [
        "- Check out of the hotel.",
        "- Transfer to Airport / End of service."
    ])
]

for title, rows in sections:
    pdf.set_font('helvetica', 'B', 12)
    pdf.cell(0, 10, title, border=0, new_x='LMARGIN', new_y='NEXT')
    pdf.set_font('helvetica', '', 11)
    for row in rows:
        if row.startswith('('):
            pdf.set_font('helvetica', 'I', 10)
            pdf.multi_cell(0, 6, row, new_x='LMARGIN', new_y='NEXT')
            pdf.set_font('helvetica', '', 11)
        else:
            pdf.multi_cell(0, 6, row, align='L', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(3)

pdf.add_page()
pdf.set_font('helvetica', 'B', 16)
pdf.cell(0, 10, 'QUOTATION', align='C', border=0, new_x='LMARGIN', new_y='NEXT')
pdf.set_font('helvetica', 'B', 14)
pdf.cell(0, 10, 'YIDA Co. Ltd', align='C', border=0, new_x='LMARGIN', new_y='NEXT')
pdf.set_font('helvetica', 'I', 11)
pdf.cell(0, 10, 'Osaka, Japan', align='C', border=0, new_x='LMARGIN', new_y='NEXT')
pdf.ln(5)

pdf.set_font('helvetica', 'B', 12)
pdf.cell(0, 10, '2026 Beginning of June - Beginning of Sept Quotation', border=0, new_x='LMARGIN', new_y='NEXT')
pdf.set_font('helvetica', '', 11)
pdf.multi_cell(0, 6, "- 4 Pax: 326,000 JPY per person (7 seats Alfa)\n- 6 Pax: 274,000 JPY per person (10 seats Hiace)\n- 8 Pax: 248,000 JPY per person (10 seats Hiace)", align='L', new_x='LMARGIN', new_y='NEXT')
pdf.set_font('helvetica', 'I', 10)
pdf.multi_cell(0, 6, "(Single supplement: 15,000 JPY per night, per person)", align='L', new_x='LMARGIN', new_y='NEXT')
pdf.ln(5)

pdf.set_font('helvetica', 'B', 12)
pdf.cell(0, 10, 'Price Includes:', border=0, new_x='LMARGIN', new_y='NEXT')
pdf.set_font('helvetica', '', 11)
pdf.multi_cell(0, 6, "- 3 hotel nights with breakfast (Tokyo: Hotel JAL City Tokyo Toyosu, Mishima Onsen: Fugaku no Yu).\n- 1 night at the 6th station mountain hut.\n- 4 days bus usage + 2 transfers limit to 10 hours and 300km per day.\n- Meals: 4x Buffet breakfast, 3x Halal lunch (2,500 JPY), 3x Halal dinner (3,500 JPY), 1x Bento dinner.", align='L', new_x='LMARGIN', new_y='NEXT')
pdf.set_font('helvetica', 'I', 10)
pdf.multi_cell(0, 6, "(A 10% fee applies if dietary changes made after booking).", align='L', new_x='LMARGIN', new_y='NEXT')
pdf.set_font('helvetica', '', 11)
pdf.multi_cell(0, 6, "- English-speaking guide for 5 days (Malaysian speaker +8,000 JPY/day).\n- Driver and guide's accommodation.", align='L', new_x='LMARGIN', new_y='NEXT')
pdf.ln(5)

pdf.set_font('helvetica', 'B', 12)
pdf.cell(0, 10, 'Price Excludes:', border=0, new_x='LMARGIN', new_y='NEXT')
pdf.set_font('helvetica', '', 11)
pdf.multi_cell(0, 6, "- Bottled water (150 JPY per bottle on the bus).\n- Weekend & Holiday surcharges (2,000-6,000 JPY/pax/day).", align='L', new_x='LMARGIN', new_y='NEXT')

pdf.output(os.path.join(output_dir, 'Fuji_Climbing_Brochure_2026.pdf'))

print("Done generating documents.")
